"""
Export Service - handles exporting MCQ data to various formats locally.
No AI calls - all conversions done with Python logic.
"""
import json
import csv
import io
import logging
from typing import List, Dict, Any, Union

logger = logging.getLogger(__name__)
OPTION_LETTERS = ['A', 'B', 'C', 'D']

class ExportService:
    def __init__(self):
        logger.info("Export Service initialized")
    
    def load_master_json(self, file_id: str) -> List[Dict[str, Any]]:
        from backend.services.storage_service import StorageService
        storage = StorageService()
        json_content = storage.get_json_by_uuid(file_id)
        if not json_content:
            raise ValueError(f"No JSON found for file_id: {file_id}")
        try:
            mcqs = json.loads(json_content)
            logger.info(f"Loaded {len(mcqs)} MCQs from MASTER JSON")
            return mcqs
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {e}")
    
    def _ensure_four_options(self, options: List[str]) -> List[str]:
        while len(options) < 4:
            options.append('')
        return options[:4]
    
    def _get_correct_letter(self, correct_idx: int) -> str:
        return OPTION_LETTERS[correct_idx] if 0 <= correct_idx < 4 else 'A'
    
    def _filter_mcqs(self, mcqs: List[Dict[str, Any]], selected_indices: List[int] = None, removed_indices: List[int] = None) -> List[Dict[str, Any]]:
        """Filter MCQs based on selected and removed indices.
        
        Logic:
        - If selected_indices is provided, only include those indices
        - Always exclude removed_indices
        - If neither is provided, return all MCQs
        """
        if not selected_indices and not removed_indices:
            return mcqs
        
        filtered = []
        removed_set = set(removed_indices) if removed_indices else set()
        
        if selected_indices:
            # Only include selected indices that are not removed
            selected_set = set(selected_indices)
            for idx in selected_indices:
                if idx not in removed_set and idx < len(mcqs):
                    filtered.append(mcqs[idx])
        else:
            # Include all except removed
            for idx, mcq in enumerate(mcqs):
                if idx not in removed_set:
                    filtered.append(mcq)
        
        return filtered
    
    # JSON
    def export_json(self, mcqs: List[Dict[str, Any]], pretty: bool = True) -> str:
        return json.dumps(mcqs, indent=2, ensure_ascii=False) if pretty else json.dumps(mcqs, ensure_ascii=False)
    
    # CSV
    def export_csv(self, mcqs: List[Dict[str, Any]]) -> str:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['ID', 'Question', 'Option A', 'Option B', 'Option C', 'Option D', 'Correct Answer'])
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            writer.writerow([mcq.get('id', ''), mcq.get('question', ''), options[0], options[1], options[2], options[3], self._get_correct_letter(correct_idx)])
        return output.getvalue()
    
    # TXT
    def export_txt(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = []
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(f"Q{mcq.get('id', '')}: {mcq.get('question', '')}")
            for idx, option in enumerate(options):
                marker = " [CORRECT]" if idx == correct_idx else ""
                lines.append(f"   {OPTION_LETTERS[idx]}) {option}{marker}")
            lines.append("")
        return '\n'.join(lines)
    
    # Markdown
    def export_markdown(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = ["# MCQ Questions", ""]
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(f"## Question {mcq.get('id', '')}\n**{mcq.get('question', '')}**\n")
            for idx, option in enumerate(options):
                lines.append(f"- **{OPTION_LETTERS[idx]}) {option}** ✅" if idx == correct_idx else f"- {OPTION_LETTERS[idx]}) {option}")
            lines.append(f"\n*Answer: {self._get_correct_letter(correct_idx)}*\n---\n")
        return '\n'.join(lines)
    
    # HTML
    def export_html(self, mcqs: List[Dict[str, Any]]) -> str:
        html = "<!DOCTYPE html><html><head><meta charset='UTF-8'><title>MCQ Questions</title></head><body><h1>MCQ Questions</h1>"
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            html += f"<div class='mcq'><h3>Q{mcq.get('id', '')}: {mcq.get('question', '')}</h3><ul>"
            for idx, option in enumerate(options):
                html += f"<li>{OPTION_LETTERS[idx]}) {option}</li>"
            html += f"</ul><p><strong>Answer: {self._get_correct_letter(correct_idx)}</strong></p></div>"
        return html + "</body></html>"
    
    # Excel (XLSX)
    def export_excel(self, mcqs: List[Dict[str, Any]]) -> bytes:
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            from io import BytesIO
            wb = Workbook()
            ws = wb.active
            ws.title = "MCQ Questions"
            headers = ['ID', 'Question', 'Option A', 'Option B', 'Option C', 'Option D', 'Correct Answer']
            for col, header in enumerate(headers, 1):
                ws.cell(row=1, column=col, value=header)
            
            # Calculate dynamic column widths based on content
            col_widths = [0] * 7  # 7 columns
            
            # Check header widths
            for col, header in enumerate(headers):
                col_widths[col] = len(str(header)) + 2
            
            # Check data widths
            for row_idx, mcq in enumerate(mcqs, 2):
                options = self._ensure_four_options(mcq.get('options', []))
                correct_idx = mcq.get('correct_answer', 0)
                
                values = [mcq.get('id', ''), mcq.get('question', '')] + options + [self._get_correct_letter(correct_idx)]
                for col_idx, val in enumerate(values):
                    val_str = str(val) if val else ''
                    # Set minimum width of 10, maximum width of 50
                    col_widths[col_idx] = max(col_widths[col_idx], min(len(val_str) + 2, 50))
            
            # Apply column widths (openpyxl uses character width)
            for col_idx, width in enumerate(col_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width
            
            # Fill data
            for row_idx, mcq in enumerate(mcqs, 2):
                options = self._ensure_four_options(mcq.get('options', []))
                correct_idx = mcq.get('correct_answer', 0)
                for col_idx, val in enumerate([mcq.get('id', ''), mcq.get('question', '')] + options + [self._get_correct_letter(correct_idx)], 1):
                    ws.cell(row=row_idx, column=col_idx, value=val)
            
            output = BytesIO()
            wb.save(output)
            return output.getvalue()
        except ImportError:
            raise ValueError("Excel export requires openpyxl. Install with: pip install openpyxl")
    
    # XML
    def export_xml(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = ['<?xml version="1.0" encoding="UTF-8"?><mcqs>']
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(f"<mcq><id>{mcq.get('id', '')}</id><question>{mcq.get('question', '')}</question><options>")
            for idx, opt in enumerate(options):
                is_correct = 'true' if idx == correct_idx else 'false'
                lines.append(f"<option letter='{OPTION_LETTERS[idx]}' correct='{is_correct}'>{opt}</option>")
            lines.append("</options></mcq>")
        lines.append("</mcqs>")
        return '\n'.join(lines)
    
    # YAML
    def export_yaml(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = ["# MCQ Questions Export", ""]
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(f"- id: {mcq.get('id', '')}")
            lines.append(f'  question: "{mcq.get("question", "")}"')
            lines.append("  options:")
            lines.append(f'    A: "{options[0]}"')
            lines.append(f'    B: "{options[1]}"')
            lines.append(f'    C: "{options[2]}"')
            lines.append(f'    D: "{options[3]}"')
            lines.append(f"  answer: {self._get_correct_letter(correct_idx)}")
            lines.append("")
        return '\n'.join(lines)
    
    # SQL Insert Script
    def export_sql(self, mcqs: List[Dict[str, Any]], table_name: str = 'mcqs') -> str:
        lines = [
            "-- MCQ Questions SQL Insert Statements",
            "-- Generated by MCQ Extractor AI",
            "",
            f"CREATE TABLE IF NOT EXISTS {table_name} (",
            "    id INTEGER PRIMARY KEY,",
            "    question TEXT NOT NULL,",
            "    option_a TEXT NOT NULL,",
            "    option_b TEXT NOT NULL,",
            "    option_c TEXT NOT NULL,",
            "    option_d TEXT NOT NULL,",
            "    correct_answer CHAR(1) NOT NULL",
            ");",
            ""
        ]
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            q = mcq.get('question', '').replace("'", "''")
            opts = [opt.replace("'", "''") for opt in options]
            lines.append(f"INSERT INTO {table_name} (id, question, option_a, option_b, option_c, option_d, correct_answer) VALUES ({mcq.get('id', 0)}, '{q}', '{opts[0]}', '{opts[1]}', '{opts[2]}', '{opts[3]}', '{self._get_correct_letter(correct_idx)}');")
        return '\n'.join(lines)
    
    # Aiken Format
    def export_aiken(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = []
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(mcq.get('question', ''))
            for idx, option in enumerate(options):
                lines.append(f"{OPTION_LETTERS[idx]}) {option}")
            lines.append(f"ANSWER: {self._get_correct_letter(correct_idx)}")
            lines.append("")
        return '\n'.join(lines)
    
    # GIFT Format
    def export_gift(self, mcqs: List[Dict[str, Any]]) -> str:
        lines = ["// GIFT Format Export", "// Generated by MCQ Extractor AI", ""]
        for mcq in mcqs:
            options = self._ensure_four_options(mcq.get('options', []))
            correct_idx = mcq.get('correct_answer', 0)
            lines.append(f"::{mcq.get('id', '')}::{mcq.get('question', '')}{{")
            for idx, option in enumerate(options):
                prefix = '=' if idx == correct_idx else '~'
                lines.append(f"    {prefix}{option}")
            lines.append("}")
            lines.append("")
        return '\n'.join(lines)
    
    # ==================== NEW PRINT-READY FORMATS ====================
    
    # Question Paper PDF (questions + MCQs only, no answers)
    def export_question_pdf(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Export question paper as PDF without answers."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                                   leftMargin=20*mm, rightMargin=20*mm)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18,
                                        alignment=TA_CENTER, spaceAfter=20)
            question_style = ParagraphStyle('Question', parent=styles['Normal'], fontSize=11,
                                           spaceBefore=10, spaceAfter=5)
            option_style = ParagraphStyle('Option', parent=styles['Normal'], fontSize=10,
                                          leftIndent=20)
            
            story = []
            story.append(Paragraph("Question Paper", title_style))
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"Total Questions: {len(mcqs)}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            for idx, mcq in enumerate(mcqs, 1):
                options = self._ensure_four_options(mcq.get('options', []))
                q_text = f"<b>Q{idx}.</b> {mcq.get('question', '')}"
                story.append(Paragraph(q_text, question_style))
                
                for opt_idx, option in enumerate(options):
                    opt_text = f"{OPTION_LETTERS[opt_idx]}) {option}"
                    story.append(Paragraph(opt_text, option_style))
                
                story.append(Spacer(1, 10))
                
                if idx % 25 == 0:
                    story.append(PageBreak())
            
            doc.build(story)
            return buffer.getvalue()
        except ImportError:
            raise ValueError("PDF export requires reportlab. Install with: pip install reportlab")
    
    # Answer Key PDF (only answers with question IDs)
    def export_answer_key_pdf(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Export answer key as PDF."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.enums import TA_CENTER
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                                   leftMargin=20*mm, rightMargin=20*mm)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18,
                                        alignment=TA_CENTER, spaceAfter=20)
            
            story = []
            story.append(Paragraph("Answer Key", title_style))
            story.append(Spacer(1, 20))
            
            data = [['Q.No.', 'Answer']]
            for mcq in mcqs:
                q_id = mcq.get('id', '')
                correct_idx = mcq.get('correct_answer', 0)
                answer = self._get_correct_letter(correct_idx)
                data.append([str(q_id), answer])
            
            table = Table(data, colWidths=[50*mm, 30*mm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
            ]))
            
            story.append(table)
            doc.build(story)
            return buffer.getvalue()
        except ImportError:
            raise ValueError("PDF export requires reportlab. Install with: pip install reportlab")
    
    # OMR Sheet PDF - Improved layout with bubbles
    def export_omr_separate_pdf(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Export OMR Answer Sheet with properly aligned bubbles.
        
        Features:
        - Section 1: Question Paper with all MCQs
        - Section 2: OMR Answer Sheet on new page
        - Perfect circular bubbles (12px diameter)
        - 4-column layout (Q1-25, Q26-50, Q51-75, Q76-100 per page)
        - A4 page with 40px margins
        - Header with student info fields
        - Instructions section
        - Dynamic adaptation to question count
        - Memory efficient implementation
        """
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER
            from reportlab.pdfgen import canvas
            from io import BytesIO
            
            # Configuration
            PAGE_WIDTH, PAGE_HEIGHT = A4
            MARGIN = 40*mm
            FONT_SIZE = 11
            LINE_SPACING = 1.2
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4, 
                topMargin=MARGIN, 
                bottomMargin=MARGIN,
                leftMargin=MARGIN, 
                rightMargin=MARGIN
            )
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=16, alignment=TA_CENTER, spaceAfter=20)
            question_style = ParagraphStyle('Question', parent=styles['Normal'], fontSize=FONT_SIZE, spaceBefore=12, spaceAfter=6)
            option_style = ParagraphStyle('Option', parent=styles['Normal'], fontSize=FONT_SIZE, leftIndent=15, spaceAfter=2)
            
            story = []
            
            # ===== SECTION 1: QUESTION PAPER =====
            story.append(Paragraph("Question Paper", title_style))
            story.append(Spacer(1, 10))
            story.append(Paragraph(f"Total Questions: {len(mcqs)}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            for idx, mcq in enumerate(mcqs, 1):
                options = self._ensure_four_options(mcq.get('options', []))
                q_text = f"<b>Q{idx}.</b> {mcq.get('question', '')}"
                story.append(Paragraph(q_text, question_style))
                
                for opt_idx, option in enumerate(options):
                    opt_text = f"{OPTION_LETTERS[opt_idx]}) {option}"
                    story.append(Paragraph(opt_text, option_style))
                
                story.append(Spacer(1, 8))
                
                # Page break after every 25 questions
                if idx % 25 == 0:
                    story.append(PageBreak())
            
            # ===== SECTION 2: OMR ANSWER SHEET (New Page) =====
            # Generate combined PDF with question paper and OMR sheet
            return self._generate_combined_pdf(mcqs)
            
        except ImportError:
            raise ValueError("PDF export requires reportlab. Install with: pip install reportlab")
    
    def _generate_combined_pdf(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Generate combined PDF with Question Paper and OMR Sheet."""
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.pdfgen import canvas
        from io import BytesIO
        
        PAGE_WIDTH, PAGE_HEIGHT = A4
        MARGIN = 40*mm
        FONT_SIZE = 11
        
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)
        
        # ===== SECTION 1: QUESTION PAPER =====
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(PAGE_WIDTH/2, PAGE_HEIGHT - 50, "Question Paper")
        
        c.setFont("Helvetica", 10)
        c.drawCentredString(PAGE_WIDTH/2, PAGE_HEIGHT - 70, f"Total Questions: {len(mcqs)}")
        
        y_position = PAGE_HEIGHT - 110
        line_height = 18
        
        for idx, mcq in enumerate(mcqs, 1):
            # Check if we need a new page
            if y_position < MARGIN + 50:
                c.showPage()
                c.setFont("Helvetica", FONT_SIZE)
                y_position = PAGE_HEIGHT - MARGIN
            
            # Question
            c.setFont("Helvetica-Bold", FONT_SIZE)
            question_text = f"Q{idx}. {mcq.get('question', '')}"
            
            # Simple text wrapping
            lines = self._wrap_text(question_text, PAGE_WIDTH - 2*MARGIN - 20, c)
            for line in lines:
                c.drawString(MARGIN, y_position, line)
                y_position -= line_height
            
            # Options
            c.setFont("Helvetica", FONT_SIZE)
            options = self._ensure_four_options(mcq.get('options', []))
            for opt_idx, option in enumerate(options):
                opt_text = f"{OPTION_LETTERS[opt_idx]}) {option}"
                opt_lines = self._wrap_text(opt_text, PAGE_WIDTH - 2*MARGIN - 40, c)
                for line in opt_lines:
                    c.drawString(MARGIN + 15, y_position, line)
                    y_position -= 14
            
            y_position -= 15  # Extra spacing between questions
        
        # ===== SECTION 2: OMR ANSWER SHEET (New Page) =====
        self._draw_omr_sheet(c, len(mcqs), PAGE_WIDTH, PAGE_HEIGHT, MARGIN)
        
        c.save()
        return buffer.getvalue()
    
    def _wrap_text(self, text: str, max_width: float, c) -> list:
        """Simple text wrapping."""
        lines = []
        words = text.split()
        current_line = ""
        
        for word in words:
            test_line = current_line + " " + word if current_line else word
            if c.stringWidth(test_line, c._fontname, c._fontsize) <= max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return lines if lines else [text]
    
    def _draw_omr_sheet(self, c, num_questions: int, page_width: float, page_height: float, margin: float):
        """Draw the OMR Answer Sheet section."""
        from reportlab.lib import colors
        from reportlab.lib.units import mm
        
        c.showPage()
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(page_width/2, page_height - 50, "OMR ANSWER SHEET")
        
        # Student Info Fields
        c.setFont("Helvetica", 10)
        c.drawString(margin, page_height - 90, "Exam Name:")
        c.line(margin + 30*mm, page_height - 90, margin + 80*mm, page_height - 90)
        
        c.drawString(margin + 85*mm, page_height - 90, "Student Name:")
        c.line(margin + 120*mm, page_height - 90, margin + 170*mm, page_height - 90)
        
        c.drawString(margin, page_height - 110, "Roll Number:")
        c.line(margin + 30*mm, page_height - 110, margin + 70*mm, page_height - 110)
        
        c.drawString(margin + 75*mm, page_height - 110, "Date:")
        c.line(margin + 95*mm, page_height - 110, margin + 125*mm, page_height - 110)
        
        # Instructions
        c.setFont("Helvetica-Bold", 10)
        c.drawString(margin, page_height - 140, "Instructions:")
        c.setFont("Helvetica", 9)
        c.drawString(margin + 5, page_height - 155, "• Fill the correct bubble completely")
        c.drawString(margin + 5, page_height - 170, "• Use blue or black pen only")
        c.drawString(margin + 5, page_height - 185, "• Do not tick, cross, or mark outside bubbles")
        
        # Separator line
        c.setLineWidth(1)
        c.line(margin, page_height - 200, page_width - margin, page_height - 200)
        
        # OMR Bubbles
        self._draw_omr_bubbles_grid(c, num_questions, page_width, page_height, margin)
    
    def _draw_omr_bubbles_grid(self, c, num_questions: int, page_width: float, page_height: float, margin: float):
        """Draw OMR bubbles in 2-column grid layout with proper spacing."""
        from reportlab.lib import colors
        
        # Bubble specifications
        BUBBLE_RADIUS = 5       # 10px diameter
        BUBBLE_SPACING = 30     # Increased spacing between bubbles
        ROW_SPACING = 32        # Increased row spacing
        BORDER_WIDTH = 1.5      # 1.5px border
        
        # Max questions per column - exactly 19 per column
        MAX_QUESTIONS_PER_COL = 19
        COLS = 2                # 2 columns per page
        
        # Calculate column width
        col_width = (page_width - 2*margin) / COLS
        
        # Content area
        content_top = page_height - 220
        
        # Calculate number of pages needed
        questions_per_page = MAX_QUESTIONS_PER_COL * COLS
        num_pages = (num_questions + questions_per_page - 1) // questions_per_page
        if num_pages == 0 and num_questions > 0:
            num_pages = 1
        if num_questions == 0:
            num_pages = 1
        
        # Track question number across pages
        current_q = 1
        
        for page_idx in range(num_pages):
            if page_idx > 0:
                c.showPage()
                # Redraw header on new page
                c.setFont("Helvetica-Bold", 16)
                c.drawCentredString(page_width/2, page_height - 50, "OMR ANSWER SHEET (Continued)")
                c.setLineWidth(1)
                c.line(margin, page_height - 70, page_width - margin, page_height - 70)
                content_top = page_height - 100
            
            # Draw column headers
            c.setFont("Helvetica-Bold", 9)
            
            # Calculate how many questions go in each column on this page
            # Column 1: fill up to MAX_QUESTIONS_PER_COL or remaining
            col1_q_count = min(MAX_QUESTIONS_PER_COL, num_questions - (current_q - 1))
            col1_start = current_q
            col1_end = current_q + col1_q_count - 1
            
            # Column 2: starts right after column 1 ends
            col2_start = col1_end + 1
            col2_q_count = min(MAX_QUESTIONS_PER_COL, num_questions - col1_q_count - (current_q - 1))
            col2_end = col2_start + col2_q_count - 1
            
            # Draw column headers
            if col1_q_count > 0:
                col1_label = f"Q{col1_start}-{col1_end}"
                c.drawCentredString(margin + col_width/2, content_top, col1_label)
            
            if col2_q_count > 0:
                col2_label = f"Q{col2_start}-{col2_end}"
                c.drawCentredString(margin + col_width * 1.5, content_top, col2_label)
            
            # Draw bubbles for column 1
            bubble_y_start = content_top - 30
            
            if col1_q_count > 0:
                for row_idx in range(col1_q_count):
                    q_num = col1_start + row_idx
                    col_left = margin  # Column 1
                    row_y = bubble_y_start - (row_idx * ROW_SPACING)
                    
                    # Draw question number
                    c.setFont("Helvetica", 9)
                    c.drawRightString(col_left + 20, row_y, str(q_num))
                    
                    # Draw option bubbles (A, B, C, D)
                    bubble_start_x = col_left + 30
                    
                    for opt_idx, opt_label in enumerate(['A', 'B', 'C', 'D']):
                        c.setFont("Helvetica", 9)
                        c.drawString(bubble_start_x + opt_idx * (BUBBLE_RADIUS * 2 + BUBBLE_SPACING), row_y, opt_label)
                        
                        bubble_x = bubble_start_x + opt_idx * (BUBBLE_RADIUS * 2 + BUBBLE_SPACING) + 15
                        
                        c.setStrokeColor(colors.black)
                        c.setLineWidth(BORDER_WIDTH)
                        c.circle(bubble_x, row_y, BUBBLE_RADIUS, stroke=1, fill=0)
            
            # Draw bubbles for column 2
            if col2_q_count > 0:
                for row_idx in range(col2_q_count):
                    q_num = col2_start + row_idx
                    col_left = margin + col_width  # Column 2
                    row_y = bubble_y_start - (row_idx * ROW_SPACING)
                    
                    # Draw question number
                    c.setFont("Helvetica", 9)
                    c.drawRightString(col_left + 20, row_y, str(q_num))
                    
                    # Draw option bubbles (A, B, C, D)
                    bubble_start_x = col_left + 30
                    
                    for opt_idx, opt_label in enumerate(['A', 'B', 'C', 'D']):
                        c.setFont("Helvetica", 9)
                        c.drawString(bubble_start_x + opt_idx * (BUBBLE_RADIUS * 2 + BUBBLE_SPACING), row_y, opt_label)
                        
                        bubble_x = bubble_start_x + opt_idx * (BUBBLE_RADIUS * 2 + BUBBLE_SPACING) + 15
                        
                        c.setStrokeColor(colors.black)
                        c.setLineWidth(BORDER_WIDTH)
                        c.circle(bubble_x, row_y, BUBBLE_RADIUS, stroke=1, fill=0)
            
            # Update current_q for next page
            current_q = col2_end + 1
        
        # Footer
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        c.drawCentredString(page_width/2, 20, f"Page 2 of {num_pages + 1}")
        c.setFillColor(colors.black)
    
    def _draw_omr_header(self, c, page_width, page_height, margin, colors, mm):
        """Draw the OMR sheet header with student information fields."""
        c.setStrokeColor(colors.black)
        c.setLineWidth(0.5)
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(page_width / 2, page_height - 35*mm, "OMR ANSWER SHEET")
        
        # Exam Name field
        c.setFont("Helvetica", 10)
        c.drawString(margin, page_height - 50*mm, "Exam Name:")
        c.setLineWidth(0.3)
        c.line(margin + 25*mm, page_height - 50*mm, margin + 85*mm, page_height - 50*mm)
        
        # Student Name field
        c.drawString(margin + 95*mm, page_height - 50*mm, "Student Name:")
        c.line(margin + 125*mm, page_height - 50*mm, margin + 185*mm, page_height - 50*mm)
        
        # Roll Number field
        c.drawString(margin, page_height - 58*mm, "Roll Number:")
        c.line(margin + 25*mm, page_height - 58*mm, margin + 65*mm, page_height - 58*mm)
        
        # Date field
        c.drawString(margin + 75*mm, page_height - 58*mm, "Date:")
        c.line(margin + 95*mm, page_height - 58*mm, margin + 125*mm, page_height - 58*mm)
        
        c.setLineWidth(0.5)
    
    def _draw_omr_instructions(self, c, page_width, margin, page_height, colors, mm):
        """Draw the instructions section."""
        c.setFont("Helvetica-Bold", 9)
        c.setFillColor(colors.black)
        c.drawString(margin, page_height - 75*mm, "Instructions:")
        
        c.setFont("Helvetica", 8)
        c.drawString(margin + 5*mm, page_height - 83*mm, "• Fill the correct bubble completely")
        c.drawString(margin + 5*mm, page_height - 89*mm, "• Use blue or black pen only")
        c.drawString(margin + 5*mm, page_height - 95*mm, "• Do not tick, cross, or mark outside bubbles")
        
        # Draw separator line
        c.setLineWidth(1)
        c.line(margin, page_height - 102*mm, page_width - margin, page_height - 102*mm)
        c.setLineWidth(0.5)
    
    def _draw_omr_bubbles(self, c, start_q, end_q, questions_per_col, margin, page_width, 
                          page_height, bubble_dia, bubble_spacing, colors, mm):
        """Draw OMR bubbles in 4-column layout."""
        content_top = page_height - 110*mm  # Start below header/instructions
        content_bottom = 40*mm  # Bottom margin
        
        # Column configuration (4 columns)
        num_cols = 4
        col_width = (page_width - 2*margin) / num_cols
        
        # Row configuration
        row_height = 8*mm  # Height for each question row
        num_rows = questions_per_col  # 25 questions per column
        
        # Draw column headers (Q range)
        c.setFont("Helvetica-Bold", 9)
        for col_idx in range(num_cols):
            col_left = margin + col_idx * col_width
            q_start = (col_idx * questions_per_col) + 1
            q_end = min((col_idx + 1) * questions_per_col, end_q - 1)
            if q_start <= q_end:
                col_label = f"Q{q_start}-{q_end}"
                c.drawCentredString(col_left + col_width/2, content_top + 3*mm, col_label)
        
        # Draw bubbles
        c.setFont("Helvetica", 8)
        bubble_radius = bubble_dia / 2
        
        for q_num in range(start_q, end_q):
            # Calculate column and row position
            global_row = q_num - start_q
            col_idx = global_row // questions_per_col
            row_idx = global_row % questions_per_col
            
            if col_idx >= num_cols:
                break
            
            col_left = margin + col_idx * col_width
            row_y = content_top - 10*mm - (row_idx * row_height)
            
            # Question number
            q_str = str(q_num)
            c.setFont("Helvetica", 8)
            c.drawRightString(col_left + 10*mm, row_y - 3*mm, q_str)
            
            # Draw option bubbles (A, B, C, D)
            bubble_start_x = col_left + 12*mm
            option_labels = ['A', 'B', 'C', 'D']
            
            for opt_idx, opt_label in enumerate(option_labels):
                bubble_x = bubble_start_x + opt_idx * (bubble_dia + bubble_spacing)
                bubble_y = row_y - bubble_radius
                
                # Draw bubble outline (circle)
                c.setStrokeColor(colors.black)
                c.setLineWidth(0.8)  # Thicker line for printing
                c.circle(bubble_x, bubble_y, bubble_radius, stroke=1, fill=0)
                
                # Draw option label below bubble
                c.setFont("Helvetica", 7)
                c.drawCentredString(bubble_x, bubble_y - bubble_radius - 3*mm, opt_label)
    
    def _draw_omr_footer(self, c, page_width, page_height, margin, current_page, total_pages, colors, mm):
        """Draw the OMR sheet footer."""
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.grey)
        footer_text = f"Page {current_page} of {total_pages}"
        c.drawCentredString(page_width / 2, 20*mm, footer_text)
        c.setFillColor(colors.black)
    
    # Tabular PDF (6 columns: Q#, Question, Opt1-4, Answer)
    def export_tabular_pdf(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Export MCQs in tabular format PDF with text wrapping."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from io import BytesIO
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), topMargin=10*mm, 
                                   bottomMargin=10*mm, leftMargin=10*mm, rightMargin=10*mm)
            
            # Create styles for wrapped text
            styles = getSampleStyleSheet()
            cell_style = ParagraphStyle('Cell', parent=styles['Normal'], fontSize=7, 
                                        leading=9, alignment=TA_LEFT)
            header_style = ParagraphStyle('Header', parent=styles['Normal'], fontSize=8, 
                                          leading=10, alignment=TA_CENTER)
            
            # Build table data with Paragraph for text wrapping
            data = [[Paragraph('Q.No.', header_style), 
                     Paragraph('Question', header_style), 
                     Paragraph('Option A', header_style), 
                     Paragraph('Option B', header_style), 
                     Paragraph('Option C', header_style), 
                     Paragraph('Option D', header_style), 
                     Paragraph('Answer', header_style)]]
            
            for mcq in mcqs:
                q_id = str(mcq.get('id', ''))
                question = mcq.get('question', '')
                options = self._ensure_four_options(mcq.get('options', []))
                correct_idx = mcq.get('correct_answer', 0)
                answer = self._get_correct_letter(correct_idx)
                
                data.append([
                    Paragraph(q_id, cell_style),
                    Paragraph(question, cell_style),
                    Paragraph(options[0] if options[0] else '', cell_style),
                    Paragraph(options[1] if options[1] else '', cell_style),
                    Paragraph(options[2] if options[2] else '', cell_style),
                    Paragraph(options[3] if options[3] else '', cell_style),
                    Paragraph(answer, cell_style)
                ])
            
            # Dynamic column widths - wider for Question and Options
            table = Table(data, colWidths=[15*mm, 65*mm, 45*mm, 45*mm, 45*mm, 45*mm, 15*mm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 3),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))
            
            doc.build([table])
            return buffer.getvalue()
        except ImportError:
            raise ValueError("PDF export requires reportlab. Install with: pip install reportlab")
    
    # DOCX Question Paper
    def export_docx(self, mcqs: List[Dict[str, Any]]) -> bytes:
        """Export question paper as DOCX."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from io import BytesIO
            
            doc = Document()
            
            title = doc.add_heading('Question Paper', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f'Total Questions: {len(mcqs)}')
            doc.add_paragraph()
            
            for idx, mcq in enumerate(mcqs, 1):
                options = self._ensure_four_options(mcq.get('options', []))
                
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f'Q{idx}. {mcq.get("question", "")}')
                q_run.bold = True
                
                for opt_idx, option in enumerate(options):
                    opt_para = doc.add_paragraph(f'   {OPTION_LETTERS[opt_idx]}) {option}')
                
                doc.add_paragraph()
            
            buffer = BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        except ImportError:
            raise ValueError("DOCX export requires python-docx. Install with: pip install python-docx")
    
    # Main export method
    def export(self, file_id: str, format: str, selected_indices: List[int] = None, removed_indices: List[int] = None) -> Union[str, bytes]:
        mcqs = self.load_master_json(file_id)
        if not mcqs:
            raise ValueError("No MCQs found in the file")
        
        # Filter MCQs based on selected and removed indices
        mcqs = self._filter_mcqs(mcqs, selected_indices, removed_indices)
        if not mcqs:
            raise ValueError("No MCQs available after filtering. Please select at least one question.")
        
        fmt = format.lower()
        
        if fmt == 'json':
            return self.export_json(mcqs)
        elif fmt == 'csv':
            return self.export_csv(mcqs)
        elif fmt == 'txt':
            return self.export_txt(mcqs)
        elif fmt == 'markdown':
            return self.export_markdown(mcqs)
        elif fmt == 'html':
            return self.export_html(mcqs)
        elif fmt == 'xml':
            return self.export_xml(mcqs)
        elif fmt == 'yaml':
            return self.export_yaml(mcqs)
        elif fmt == 'sql':
            return self.export_sql(mcqs)
        elif fmt == 'aiken':
            return self.export_aiken(mcqs)
        elif fmt == 'gift':
            return self.export_gift(mcqs)
        elif fmt == 'excel':
            return self.export_excel(mcqs)
        # New print-ready formats
        elif fmt == 'question_pdf':
            return self.export_question_pdf(mcqs)
        elif fmt == 'answer_key_pdf':
            return self.export_answer_key_pdf(mcqs)
        elif fmt == 'omr_pdf':
            return self.export_omr_separate_pdf(mcqs)
        elif fmt == 'tabular_pdf':
            return self.export_tabular_pdf(mcqs)
        elif fmt == 'docx':
            return self.export_docx(mcqs)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def get_supported_formats(self) -> List[str]:
        return ['json', 'csv', 'txt', 'markdown', 'html', 'xml', 'yaml', 'sql', 'aiken', 'gift', 'excel',
                'question_pdf', 'answer_key_pdf', 'omr_pdf', 'tabular_pdf', 'docx']

